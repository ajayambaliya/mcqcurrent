import io
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from datetime import datetime
import pymongo
from deep_translator import GoogleTranslator
import asyncio
import telegram
import tempfile
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml import OxmlElement
from PIL import Image  # Pillow for image conversion

# MongoDB setup with defaults from environment variables (GitHub secrets)
DB_NAME = os.environ.get('DB_NAME')
COLLECTION_NAME = os.environ.get('COLLECTION_NAME')
MONGO_CONNECTION_STRING = os.environ.get('MONGO_CONNECTION_STRING')

# Initialize MongoDB client and collection if credentials are available
client = None
db = None
collection = None

# Validate MongoDB credentials if not in testing mode
if not os.environ.get('TESTING_MODE') and not all([DB_NAME, COLLECTION_NAME, MONGO_CONNECTION_STRING]):
    raise ValueError("Missing MongoDB environment variables: DB_NAME, COLLECTION_NAME, or MONGO_CONNECTION_STRING")
elif all([DB_NAME, COLLECTION_NAME, MONGO_CONNECTION_STRING]):
    try:
        client = pymongo.MongoClient(MONGO_CONNECTION_STRING)
        db = client[DB_NAME]
        collection = db[COLLECTION_NAME]
    except pymongo.errors.ConnectionError as e:
        print(f"Failed to connect to MongoDB: {str(e)}")
        # Don't raise an error here to allow testing without MongoDB

def fetch_article_urls(base_url, pages):
    article_urls = []
    for page in range(1, pages + 1):
        url = base_url if page == 1 else f"{base_url}page/{page}/"
        try:
            print(f"Fetching URLs from: {url}")
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Find article URLs using the structure provided by the user
            post_data_divs = soup.find_all('div', class_='post-data')
            for div in post_data_divs:
                h3_tag = div.find('h3')
                if h3_tag:
                    a_tag = h3_tag.find('a')
                    if a_tag and a_tag.get('href'):
                        article_urls.append(a_tag['href'])
                        print(f"Found article: {a_tag.text.strip()} at {a_tag['href']}")
            
            # If no articles found with the above method, try alternative structures
            if not post_data_divs:
                # Try looking for articles
                articles = soup.find_all('article')
                for article in articles:
                    # Find the article title link
                    h2_tag = article.find('h2', class_='entry-title')
                    if h2_tag:
                        a_tag = h2_tag.find('a')
                        if a_tag and a_tag.get('href'):
                            article_urls.append(a_tag['href'])
                
                # If still no articles, try the original method
                if not articles:
                    for h1_tag in soup.find_all('h1', id='list'):
                        a_tag = h1_tag.find('a')
                        if a_tag and a_tag.get('href'):
                            article_urls.append(a_tag['href'])
        except requests.RequestException as e:
            print(f"Failed to fetch URLs from {url}: {str(e)}")
    
    # Remove duplicates while preserving order
    unique_urls = []
    for url in article_urls:
        if url not in unique_urls:
            unique_urls.append(url)
    
    print(f"Scraped {len(unique_urls)} unique URLs")
    return unique_urls

def translate_to_gujarati(text):
    try:
        translator = GoogleTranslator(source='auto', target='gu')
        return translator.translate(text)
    except Exception as e:
        print(f"Translation error: {str(e)}")
        return text

def download_and_convert_image(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        content = response.content
        if len(content) < 100:  # Check for invalid/small images
            print(f"Image at {url} is too small, likely invalid")
            return None
        
        # Convert image to PNG using Pillow
        img = Image.open(io.BytesIO(content))
        if img.mode in ('RGBA', 'P'):  # Convert transparency to white background
            img = img.convert('RGB')
        output = io.BytesIO()
        img.save(output, format='PNG')
        output.seek(0)
        return output
    except Exception as e:
        print(f"Failed to process image from {url}: {str(e)}")
        return None

async def scrape_and_get_content(url):
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Find the main content area with the new structure
        main_content = soup.find('main', id='main', class_='site-main')
        if not main_content:
            # Fallback to the old structure if new one is not found
            main_content = soup.find('div', class_='inside_post column content_width')
            if not main_content:
                raise Exception("Main content element not found")
        
        heading = main_content.find('h1', id='list')
        if not heading:
            raise Exception("Heading not found")
        
        # Extract featured image with the new structure
        image_url = None
        featured_img = main_content.find('img', class_='post-featured-image')
        if featured_img and featured_img.get('src'):
            image_url = featured_img['src']
        else:
            # Fallback to old structure
            image_div = soup.find('div', class_='featured_image')
            if image_div:
                img_tag = image_div.find('img')
                if img_tag and img_tag.get('src'):
                    image_url = img_tag['src']
        
        content_list = []
        heading_text = heading.get_text().strip()
        translated_heading = translate_to_gujarati(heading_text)
        content_list.append({'type': 'heading', 'text': translated_heading, 'image': image_url})
        content_list.append({'type': 'heading', 'text': heading_text, 'image': None})
        
        # Process only the content elements (paragraphs, headings, lists)
        # Skip breadcrumb, post-meta, comments and other non-content elements
        content_elements = []
        for tag in main_content.find_all(['p', 'h2', 'h4', 'ul', 'ol']):
            # Skip elements inside comments, sharethis, related-articles
            if tag.parent.get('id') == 'comments' or \
               (tag.parent.get('class') and any(c in ['sharethis-inline-share-buttons', 'related-articles', 'breadcrumb'] 
                                             for c in tag.parent.get('class', []))):
                continue
            content_elements.append(tag)
            
        numbered_list_counter = 1
        for tag in content_elements:
            text = tag.get_text().strip()
            if not text:
                continue
            translated_text = translate_to_gujarati(text)
            if tag.name == 'p':
                content_list.append({'type': 'paragraph', 'text': translated_text})
                content_list.append({'type': 'paragraph', 'text': text})
            elif tag.name == 'h2':
                content_list.append({'type': 'heading_2', 'text': translated_text})
                content_list.append({'type': 'heading_2', 'text': text})
            elif tag.name == 'h4':
                content_list.append({'type': 'heading_4', 'text': translated_text})
                content_list.append({'type': 'heading_4', 'text': text})
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    li_text = li.get_text().strip()
                    translated_li_text = translate_to_gujarati(li_text)
                    content_list.append({'type': 'bullet_list', 'text': translated_li_text})
                    content_list.append({'type': 'bullet_list', 'text': li_text})
            elif tag.name == 'ol':
                for li in tag.find_all('li'):
                    li_text = li.get_text().strip()
                    translated_li_text = translate_to_gujarati(li_text)
                    content_list.append({'type': 'numbered_list', 'text': translated_li_text, 'number': numbered_list_counter})
                    content_list.append({'type': 'numbered_list', 'text': li_text, 'number': numbered_list_counter})
                    numbered_list_counter += 1
        return content_list
    except Exception as e:
        print(f"Error scraping {url}: {str(e)}")
        return []

def add_paragraph_border(paragraph):
    """Add a bottom border to a paragraph using XML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8C8C8')  # Light gray
    pBdr.append(bottom)  # Fixed: Replace 'custom' with 'bottom'
    pPr.append(pBdr)

def setup_document_styles(doc):
    styles = doc.styles

    # Title Style
    title_style = styles.add_style('CoolTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 102, 204)
    title_style.font.underline = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    # Heading 1 Style
    h1_style = styles.add_style('CoolHeading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(51, 51, 51)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(8)

    # Heading 2 Style
    h2_style = styles.add_style('CoolHeading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.italic = True
    h2_style.font.color.rgb = RGBColor(0, 153, 76)
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(6)

    # Paragraph Style
    p_style = styles.add_style('CoolParagraph', WD_STYLE_TYPE.PARAGRAPH)
    p_style.font.name = 'Georgia'
    p_style.font.size = Pt(12)
    p_style.font.color.rgb = RGBColor(33, 33, 33)
    p_style.paragraph_format.line_spacing = 1.15
    p_style.paragraph_format.space_after = Pt(8)
    p_style.paragraph_format.first_line_indent = Inches(0.5)

    # Bullet List Style
    bullet_style = styles.add_style('CoolBulletList', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Georgia'
    bullet_style.font.size = Pt(12)
    bullet_style.font.color.rgb = RGBColor(66, 66, 66)
    bullet_style.paragraph_format.left_indent = Inches(0.75)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    bullet_style.paragraph_format.space_after = Pt(4)

    # Numbered List Style
    numbered_style = styles.add_style('CoolNumberedList', WD_STYLE_TYPE.PARAGRAPH)
    numbered_style.font.name = 'Georgia'
    numbered_style.font.size = Pt(12)
    numbered_style.font.color.rgb = RGBColor(66, 66, 66)
    numbered_style.paragraph_format.left_indent = Inches(0.75)
    numbered_style.paragraph_format.first_line_indent = Inches(-0.25)
    numbered_style.paragraph_format.space_after = Pt(4)

    # Image Paragraph Style
    img_style = styles.add_style('ImageParagraph', WD_STYLE_TYPE.PARAGRAPH)
    img_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    img_style.paragraph_format.space_before = Pt(6)
    img_style.paragraph_format.space_after = Pt(12)

def create_styled_document(content_list):
    doc = Document()
    setup_document_styles(doc)

    # Set document margins
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Add a cool title
    doc.add_paragraph(f"Current Affairs - {datetime.now().strftime('%d %B %Y')}", style='CoolTitle')
    
    for content in content_list:
        if content['type'] == 'heading':
            doc.add_paragraph(content['text'], style='CoolHeading1')
            if content.get('image'):
                image_data = download_and_convert_image(content['image'])
                if image_data:
                    p = doc.add_paragraph(style='ImageParagraph')
                    run = p.add_run()
                    run.add_picture(image_data, width=Inches(2.5), height=Inches(1.875))  # Adjusted size
        elif content['type'] == 'paragraph':
            p = doc.add_paragraph(content['text'], style='CoolParagraph')
            add_paragraph_border(p)
        elif content['type'] == 'heading_2':
            doc.add_paragraph(content['text'], style='CoolHeading2')
        elif content['type'] == 'heading_4':
            doc.add_paragraph(content['text'], style='Heading 4')
        elif content['type'] == 'bullet_list':
            doc.add_paragraph(f"• {content['text']}", style='CoolBulletList')
        elif content['type'] == 'numbered_list':
            doc.add_paragraph(f"{content['number']}. {content['text']}", style='CoolNumberedList')
        
        # Add section break after each article
        if content['type'] == 'heading' and content.get('image'):
            doc.add_section(WD_SECTION.CONTINUOUS)

    return doc

def check_and_insert_urls(urls):
    new_urls = []
    
    # If MongoDB is not available, return all URLs as new
    if collection is None:
        print("MongoDB not available. Treating all URLs as new.")
        return urls
        
    for url in urls:
        if 'daily-current-affairs-quiz' in url:
            continue
        if not collection.find_one({'url': url}):
            new_urls.append(url)
            collection.insert_one({'url': url, 'processed_date': datetime.now()})
    print(f"Found {len(new_urls)} new URLs: {new_urls}")
    return new_urls

async def send_docx_to_telegram(docx_path, bot_token, channel_id, caption):
    if not bot_token or not channel_id:
        raise ValueError("Bot token or channel ID is missing")
    
    # Convert channel_id to integer if it's a string
    if isinstance(channel_id, str):
        channel_id = int(channel_id)
    
    bot = telegram.Bot(token=bot_token)
    telegram_caption_limit = 1024
    print(f"Attempting to send to chat_id: {channel_id} (type: {type(channel_id)}) with bot token ending in: {bot_token[-6:]}")
    
    try:
        with open(docx_path, 'rb') as docx_file:
            if len(caption) > telegram_caption_limit:
                short_caption = caption[:telegram_caption_limit-3] + "..."
                await bot.send_document(
                    chat_id=channel_id,
                    document=docx_file,
                    filename=os.path.basename(docx_path),
                    caption=short_caption
                )
                await bot.send_message(chat_id=channel_id, text=caption)
            else:
                await bot.send_document(
                    chat_id=channel_id,
                    document=docx_file,
                    filename=os.path.basename(docx_path),
                    caption=caption
                )
        print("Document sent successfully to Telegram")
    except telegram.error.BadRequest as e:
        print(f"Failed to send document to Telegram: {e.message}")
        raise
    except telegram.error.TimedOut:
        print("Telegram timeout, retrying not implemented here")
        raise
    except Exception as e:
        print(f"Unexpected error sending document to Telegram: {str(e)}")
        raise

async def main():
    try:
        base_url = "https://www.gktoday.in/current-affairs/"
        article_urls = fetch_article_urls(base_url, 4)
        if not article_urls:
            print("No URLs scraped. Check website structure or connectivity.")
            return
        
        # Filter out quiz articles before checking for new URLs
        filtered_urls = [url for url in article_urls if 'quiz' not in url.lower() and 'mcq' not in url.lower()]
        print(f"Filtered out {len(article_urls) - len(filtered_urls)} quiz/MCQ articles")
        
        new_urls = check_and_insert_urls(filtered_urls)
        if not new_urls:
            print("No new URLs to process")
            return
        
        all_content = []
        english_titles = []
        for url in new_urls:
            print(f"Processing article: {url}")
            content_list = await scrape_and_get_content(url)
            if content_list:
                all_content.extend(content_list)
                # English title is at index 1 (index 0 is Gujarati title)
                english_titles.append(content_list[1]['text'])
                print(f"Added article: {content_list[1]['text']}")
            else:
                print(f"Failed to extract content from: {url}")
        
        if not all_content:
            print("No content scraped from new URLs")
            return
        
        doc = create_styled_document(all_content)
        
        current_date = datetime.now().strftime('%d-%m-%Y')
        docx_filename = f"{current_date}_Current_Affairs.docx"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            doc.save(tmp_docx.name)
            docx_path = tmp_docx.name
            print(f"Document saved to: {docx_path}")
        
        bot_token = os.environ.get('TELEGRAM_BOT_TOKEN')
        channel_id = os.environ.get('TELEGRAM_CHANNEL_ID')
        if not bot_token or not channel_id:
            raise ValueError("TELEGRAM_BOT_TOKEN or TELEGRAM_CHANNEL_ID not set in environment variables")
        
        print(f"Bot token from env ends with: {bot_token[-6:] if bot_token else 'None'}")
        print(f"Channel ID from env: {channel_id} (type: {type(channel_id)})")
        
        caption = (
            f"🎗️ {datetime.now().strftime('%d %B %Y')} Current Affairs 🎗️\n\n"
            + '\n'.join([f"👉 {title}" for title in english_titles]) + '\n\n'
        )
        
        await send_docx_to_telegram(docx_path, bot_token, channel_id, caption)
        
        os.unlink(docx_path)
        print("Temporary file deleted")
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        raise

if __name__ == "__main__":
    asyncio.run(main())
